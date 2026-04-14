"""
ADAM DNA Questionnaire Parser
Reads a filled-out ADAM DNA Questionnaire (.docx) and extracts all 13 sections
into a structured Python dictionary for downstream generators.
"""

import re
from docx import Document
from typing import Dict, List, Any, Optional


# Maps section headings to their structured keys
SECTION_MAP = {
    "1": "doctrine_identity",
    "2": "culture_graph",
    "3": "objectives_graph",
    "4": "rules_expectations",
    "5": "enterprise_memory",
    "6": "boss_scoring",
    "7": "intent_conflict",
    "8": "agentic_architecture",
    "9": "flight_recorder",
    "10": "products_services",
    "11": "temporal_regional",
    "12": "cloud_infrastructure",
    "13": "resilience_security",
}

SECTION_TITLES = {
    "1": "Doctrine Identity & Constitutional Foundation",
    "2": "CORE Engine — Culture Graph",
    "3": "CORE Engine — Objectives Graph",
    "4": "CORE Engine — Rules & Expectations Graph",
    "5": "CORE Subgraphs — Enterprise Memory",
    "6": "BOSS Scoring & Exception Economy Configuration",
    "7": "Intent Object & Doctrine Conflict Configuration",
    "8": "Agentic Architecture & Domain Configuration",
    "9": "Flight Recorder & Evidence Architecture",
    "10": "Products, Services & Operational Domain",
    "11": "Temporal & Regional Variance Configuration",
    "12": "Cloud Infrastructure Sizing & Sovereignty Architecture",
    "13": "Resilience, Idempotency & Security Posture",
}


class DNAQuestionnaireParser:
    """Parses a filled-out ADAM DNA Questionnaire .docx file."""

    def __init__(self, filepath: str):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.raw_data: Dict[str, Any] = {}
        self.company_name: str = ""
        self.company_info: Dict[str, str] = {}
        self.sections: Dict[str, Dict[str, Any]] = {}

    def parse(self) -> Dict[str, Any]:
        """Main parse entry point. Returns structured DNA data."""
        self._extract_header()
        self._extract_tables()
        self._extract_section_context()
        return self._build_dna_package()

    def _extract_header(self):
        """Extract company name and header information."""
        for para in self.doc.paragraphs[:20]:
            text = para.text.strip()
            if not text:
                continue
            # Look for company identification in early paragraphs
            if "example" not in text.lower() and "adam" not in text.lower() \
               and "version" not in text.lower() and "questionnaire" not in text.lower() \
               and "core engine" not in text.lower() and "how to use" not in text.lower() \
               and "purpose" not in text.lower() and "aligned" not in text.lower() \
               and "what this is not" not in text.lower() and "completion" not in text.lower() \
               and "autonomous" not in text.lower() and "doctrine" not in text.lower() \
               and len(text) > 2 and len(text) < 100:
                if not self.company_name:
                    self.company_name = text

    def _extract_tables(self):
        """Extract all questionnaire tables (Q&A pairs)."""
        for table in self.doc.tables:
            headers = [cell.text.strip() for cell in table.rows[0].cells]

            # Identify questionnaire tables by their header pattern
            if len(headers) >= 2 and "#" in headers[0].upper():
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) >= 3:
                        q_num = cells[0].strip()
                        question = cells[1].strip()
                        # The answer is in column 3 (or column 2 if only 2 data cols)
                        answer = cells[2].strip() if len(cells) > 2 else cells[1].strip()

                        if q_num and question:
                            section_num = q_num.split(".")[0]
                            self.raw_data[q_num] = {
                                "question_number": q_num,
                                "question": question,
                                "answer": answer,
                                "section": section_num,
                            }

    def _extract_section_context(self):
        """Extract section-level context from paragraphs (descriptions, CORE Engine Notes)."""
        current_section = None
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect section headers
            if para.style.name.startswith("Heading"):
                match = re.match(r"SECTION\s+(\d+)", text)
                if match:
                    current_section = match.group(1)
                    if current_section not in self.sections:
                        self.sections[current_section] = {
                            "title": text,
                            "context": [],
                            "core_engine_notes": [],
                        }

            elif current_section:
                if "CORE Engine Note" in text:
                    note = text.replace("CORE Engine Note:", "").strip()
                    self.sections[current_section]["core_engine_notes"].append(note)
                elif para.style.name == "Normal" and len(text) > 50:
                    self.sections[current_section]["context"].append(text)

    def _build_dna_package(self) -> Dict[str, Any]:
        """Build the complete DNA data package from parsed data."""
        dna = {
            "meta": {
                "company_name": self._get_company_name(),
                "source_file": self.filepath,
                "questionnaire_version": "1.0",
                "adam_version": "0.4",
            },
            "sections": {},
        }

        for sec_num, sec_key in SECTION_MAP.items():
            section_questions = {
                qn: qdata
                for qn, qdata in self.raw_data.items()
                if qdata["section"] == sec_num
            }
            section_context = self.sections.get(sec_num, {})

            dna["sections"][sec_key] = {
                "section_number": sec_num,
                "title": SECTION_TITLES.get(sec_num, f"Section {sec_num}"),
                "questions": section_questions,
                "context": section_context.get("context", []),
                "core_engine_notes": section_context.get("core_engine_notes", []),
            }

        # Extract key deployment parameters
        dna["deployment_params"] = self._extract_deployment_params()
        dna["governance"] = self._extract_governance_params()
        dna["boss_config"] = self._extract_boss_config()
        dna["agent_config"] = self._extract_agent_config()
        dna["infrastructure"] = self._extract_infrastructure_params()

        return dna

    def _get_company_name(self) -> str:
        """Get company name from questionnaire answers or header."""
        # Try to get from question 1.1.1 (legal name)
        if "1.1.1" in self.raw_data:
            answer = self.raw_data["1.1.1"]["answer"]
            # Extract just the company name (first part before comma or period)
            name = answer.split(",")[0].split(".")[0].strip()
            return name
        return self.company_name or "Unknown Company"

    def _extract_deployment_params(self) -> Dict[str, Any]:
        """Extract cloud infrastructure and deployment parameters from Section 12-13."""
        params = {
            "cloud_topology": {},
            "compute": {},
            "storage": {},
            "networking": {},
            "on_premises": {},
            "cross_cloud": {},
        }

        # Section 12.1 - Sovereignty-First Architecture Topology
        if "12.1.1" in self.raw_data:
            topology_answer = self.raw_data["12.1.1"]["answer"]
            params["cloud_topology"] = self._parse_topology(topology_answer)

        if "12.1.2" in self.raw_data:
            params["sovereignty_boundary"] = self.raw_data["12.1.2"]["answer"]

        if "12.1.3" in self.raw_data:
            params["data_residency"] = self.raw_data["12.1.3"]["answer"]

        # Section 12.2 - Compute
        if "12.2.1" in self.raw_data:
            params["compute"]["core_engine"] = self.raw_data["12.2.1"]["answer"]
        if "12.2.2" in self.raw_data:
            params["compute"]["agent_mesh"] = self.raw_data["12.2.2"]["answer"]
        if "12.2.3" in self.raw_data:
            params["compute"]["auto_scaling"] = self.raw_data["12.2.3"]["answer"]

        # Section 12.3 - Storage & Data
        if "12.3.1" in self.raw_data:
            params["storage"] = self.raw_data["12.3.1"]["answer"]
        if "12.3.2" in self.raw_data:
            params["networking"] = self.raw_data["12.3.2"]["answer"]

        # Section 12.4 - On-Premises
        if "12.4.1" in self.raw_data:
            params["on_premises"]["config"] = self.raw_data["12.4.1"]["answer"]
        if "12.4.2" in self.raw_data:
            params["on_premises"]["sync_strategy"] = self.raw_data["12.4.2"]["answer"]

        # Section 12.5 - Cross-Cloud
        if "12.5.1" in self.raw_data:
            params["cross_cloud"]["management"] = self.raw_data["12.5.1"]["answer"]
        if "12.5.2" in self.raw_data:
            params["cross_cloud"]["ai_services"] = self.raw_data["12.5.2"]["answer"]

        return params

    def _parse_topology(self, text: str) -> Dict[str, Any]:
        """Parse cloud topology from natural language answer."""
        topology = {
            "primary": "azure",
            "secondary": "aws",
            "on_premises": "azure_local",
            "regions": [],
        }

        text_lower = text.lower()
        if "azure" in text_lower and "primary" in text_lower:
            topology["primary"] = "azure"
        elif "aws" in text_lower and "primary" in text_lower:
            topology["primary"] = "aws"
        elif "google" in text_lower and "primary" in text_lower:
            topology["primary"] = "gcp"

        if "aws" in text_lower and ("secondary" in text_lower or "standby" in text_lower):
            topology["secondary"] = "aws"
        elif "azure" in text_lower and ("secondary" in text_lower or "standby" in text_lower):
            topology["secondary"] = "azure"

        # Extract regions
        region_patterns = [
            r"(us[- ]east|us[- ]west|eu[- ]west|southeast asia|ap[- ]northeast)",
            r"(virginia|amsterdam|tokyo|london|frankfurt|singapore|sydney|mumbai)",
        ]
        for pattern in region_patterns:
            matches = re.findall(pattern, text_lower)
            topology["regions"].extend(matches)

        return topology

    def _extract_governance_params(self) -> Dict[str, Any]:
        """Extract governance parameters from Sections 1-2."""
        governance = {
            "mission": "",
            "vision": "",
            "principles": [],
            "sacred_boundaries": [],
            "directors": {},
            "delegation_model": "",
            "values_trade_offs": [],
            "decision_philosophy": "",
        }

        if "1.1.2" in self.raw_data:
            governance["mission"] = self.raw_data["1.1.2"]["answer"]
        if "1.1.3" in self.raw_data:
            governance["vision"] = self.raw_data["1.1.3"]["answer"]
        if "1.1.4" in self.raw_data:
            governance["principles"] = self._parse_numbered_list(
                self.raw_data["1.1.4"]["answer"]
            )
        if "1.1.5" in self.raw_data:
            governance["sacred_boundaries"] = self._parse_list_items(
                self.raw_data["1.1.5"]["answer"]
            )
        if "1.2.1" in self.raw_data:
            governance["directors"] = self.raw_data["1.2.1"]["answer"]
        if "1.2.3" in self.raw_data:
            governance["delegation_model"] = self.raw_data["1.2.3"]["answer"]
        if "2.1.1" in self.raw_data:
            governance["values_trade_offs"] = self._parse_numbered_list(
                self.raw_data["2.1.1"]["answer"]
            )

        return governance

    def _extract_boss_config(self) -> Dict[str, Any]:
        """Extract BOSS scoring configuration from Section 6."""
        boss = {
            "dimensions": {},
            "routing_thresholds": {},
            "recalibration_cadence": "",
            "exception_packets": "",
            "target_exception_volume": "",
            "feedback_loop": "",
        }

        if "6.1.1" in self.raw_data:
            boss["dimensions"] = self._parse_boss_dimensions(
                self.raw_data["6.1.1"]["answer"]
            )
        if "6.1.2" in self.raw_data:
            boss["routing_thresholds"] = self._parse_boss_thresholds(
                self.raw_data["6.1.2"]["answer"]
            )
        if "6.1.3" in self.raw_data:
            boss["recalibration_cadence"] = self.raw_data["6.1.3"]["answer"]
        if "6.2.1" in self.raw_data:
            boss["exception_packets"] = self.raw_data["6.2.1"]["answer"]
        if "6.2.2" in self.raw_data:
            boss["target_exception_volume"] = self.raw_data["6.2.2"]["answer"]
        if "6.2.3" in self.raw_data:
            boss["feedback_loop"] = self.raw_data["6.2.3"]["answer"]

        return boss

    def _extract_agent_config(self) -> Dict[str, Any]:
        """Extract agentic architecture config from Section 8."""
        agents = {
            "domain_governors": {},
            "work_groups": {},
            "digital_twins": {},
        }

        for qn in ["8.1.1", "8.1.2", "8.1.3", "8.1.4", "8.1.5"]:
            if qn in self.raw_data:
                agents["domain_governors"][qn] = self.raw_data[qn]["answer"]

        if "8.2.1" in self.raw_data:
            agents["work_groups"]["corporate"] = self.raw_data["8.2.1"]["answer"]
        if "8.2.2" in self.raw_data:
            agents["work_groups"]["ai_centric"] = self.raw_data["8.2.2"]["answer"]

        if "8.3.1" in self.raw_data:
            agents["digital_twins"]["definitions"] = self.raw_data["8.3.1"]["answer"]
        if "8.3.2" in self.raw_data:
            agents["digital_twins"]["simulation_requirements"] = self.raw_data["8.3.2"]["answer"]

        return agents

    def _extract_infrastructure_params(self) -> Dict[str, Any]:
        """Extract infrastructure sizing from Section 12."""
        infra = {
            "resilience_tiers": "",
            "idempotency_requirements": "",
            "security_posture": "",
            "threat_model": "",
        }

        if "13.1" in self.raw_data:
            infra["resilience_tiers"] = self.raw_data["13.1"]["answer"]
        if "13.2" in self.raw_data:
            infra["idempotency_requirements"] = self.raw_data["13.2"]["answer"]
        if "13.3" in self.raw_data:
            infra["security_posture"] = self.raw_data["13.3"]["answer"]
        if "13.4" in self.raw_data:
            infra["threat_model"] = self.raw_data["13.4"]["answer"]

        return infra

    def _parse_numbered_list(self, text: str) -> List[str]:
        """Parse numbered items like '1) item 2) item'."""
        items = re.split(r"\d+\)\s*", text)
        return [item.strip() for item in items if item.strip()]

    def _parse_list_items(self, text: str) -> List[str]:
        """Parse items separated by 'Never' or newlines."""
        items = re.split(r"(?:Never\s+)", text)
        result = []
        for item in items:
            item = item.strip().rstrip(".")
            if item and len(item) > 5:
                if not item.startswith("Never"):
                    item = "Never " + item
                result.append(item)
        return result if result else [text]

    def _parse_boss_dimensions(self, text: str) -> Dict[str, float]:
        """Parse BOSS dimension weights from answer text."""
        dimensions = {}
        # Match patterns like "Financial Exposure: Weight 3.6" or "Security Impact: Weight 4.6"
        matches = re.findall(
            r"([\w\s&]+?):\s*Weight\s+([\d.]+)", text, re.IGNORECASE
        )
        for name, weight in matches:
            key = name.strip().lower().replace(" ", "_").replace("&", "and")
            dimensions[key] = float(weight)

        # Fallback defaults if parsing fails
        if not dimensions:
            dimensions = {
                "sovereignty_action": 4.0,
                "financial_exposure": 3.6,
                "regulatory_impact": 3.5,
                "rights_certainty": 3.2,
                "security_impact": 4.6,
                "doctrinal_alignment": 3.0,
                "reputational_risk": 3.3,
            }
        return dimensions

    def _parse_boss_thresholds(self, text: str) -> Dict[str, Dict[str, int]]:
        """Parse BOSS routing thresholds."""
        thresholds = {
            "soap": {"min": 0, "max": 10, "label": "SOAP - Safe & Optimum Performance"},
            "moderate": {"min": 11, "max": 30, "label": "MODERATE - Constrained Execution"},
            "elevated": {"min": 31, "max": 50, "label": "ELEVATED - Exception Likely"},
            "high": {"min": 51, "max": 75, "label": "HIGH - Director Review Required"},
            "ohshat": {"min": 76, "max": 100, "label": "OHSHAT - Immediate Escalation"},
        }
        return thresholds

    def get_answer(self, question_number: str) -> str:
        """Get the answer for a specific question number."""
        if question_number in self.raw_data:
            return self.raw_data[question_number]["answer"]
        return ""

    def get_section_answers(self, section_number: str) -> Dict[str, Dict]:
        """Get all Q&A pairs for a specific section."""
        return {
            qn: qdata
            for qn, qdata in self.raw_data.items()
            if qdata["section"] == section_number
        }
