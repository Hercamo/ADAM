"""
ADAM DNA Deployment Specification - Word Document Generator.
Creates a professional .docx deployment specification document per platform.
"""

import os
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from .base_generator import AGENT_CLASSES

PLATFORM_DETAILS = {
    "azure": {
        "display": "Microsoft Azure (Primary)",
        "graph_db": "Cosmos DB (Gremlin API)",
        "container": "Azure Kubernetes Service (AKS)",
        "ai": "Azure OpenAI",
        "evidence": "Azure Blob (Immutable) + Confidential Ledger",
        "vault": "Azure Key Vault (Premium HSM)",
        "messaging": "Azure Service Bus (Premium)",
        "monitoring": "Azure Monitor + Grafana Managed",
        "scoring": "Azure Data Explorer",
    },
    "aws": {
        "display": "Amazon Web Services (Warm Standby)",
        "graph_db": "Amazon Neptune (Gremlin)",
        "container": "Amazon Elastic Kubernetes Service (EKS)",
        "ai": "Amazon Bedrock",
        "evidence": "S3 (Object Lock) + QLDB",
        "vault": "AWS KMS",
        "messaging": "Amazon SQS/SNS",
        "monitoring": "CloudWatch + Managed Grafana",
        "scoring": "Amazon Timestream",
    },
    "gcp": {
        "display": "Google Cloud Platform",
        "graph_db": "Cloud Spanner",
        "container": "Google Kubernetes Engine (GKE)",
        "ai": "Vertex AI",
        "evidence": "Cloud Storage (Retention Lock)",
        "vault": "Cloud KMS",
        "messaging": "Pub/Sub",
        "monitoring": "Cloud Monitoring + Managed Grafana",
        "scoring": "BigQuery",
    },
    "kubernetes": {
        "display": "Open Source Kubernetes",
        "graph_db": "JanusGraph",
        "container": "Generic Kubernetes",
        "ai": "Ollama / vLLM (self-hosted)",
        "evidence": "MinIO (Immutable)",
        "vault": "HashiCorp Vault",
        "messaging": "NATS / RabbitMQ",
        "monitoring": "Prometheus + Grafana",
        "scoring": "TimescaleDB",
    },
    "azure-local": {
        "display": "Azure Local (On-Premises Failover)",
        "graph_db": "CORE Graph (Read-Only Cache)",
        "container": "AKS-HCI (Arc-enabled)",
        "ai": "Cached Models / Edge AI",
        "evidence": "Local Flight Recorder (sync on reconnect)",
        "vault": "Local HSM + Azure Key Vault (cached)",
        "messaging": "Local Service Bus",
        "monitoring": "Azure Arc Monitor",
        "scoring": "Cached BOSS Weights",
    },
}

class DocxSpecGenerator:
    """Generates Word document deployment specifications for each platform."""

    def __init__(self, dna_data: Dict[str, Any], output_dir: str):
        self.dna = dna_data
        self.output_dir = output_dir
        self.company_name = dna_data.get("meta", {}).get("company_name", "Company")

    def generate(self, platforms: List[str]) -> Dict[str, str]:
        """Generate .docx spec for each selected platform."""
        files = {}
        for platform in platforms:
            filepath = self._generate_platform_doc(platform)
            files[filepath] = f"Deployment Specification - {PLATFORM_DETAILS.get(platform, {}).get('display', platform)}"
        return files

    def _generate_platform_doc(self, platform: str) -> str:
        doc = Document()
        details = PLATFORM_DETAILS.get(platform, {})

        # Styles
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Title Page
        self._add_title_page(doc, platform, details)

        # Table of Contents placeholder
        doc.add_page_break()
        doc.add_heading('Table of Contents', level=1)
        doc.add_paragraph('[Auto-generated table of contents — update after opening in Word]')

        # Section 1: Executive Summary
        doc.add_page_break()
        doc.add_heading('1. Executive Summary', level=1)
        doc.add_paragraph(
            f'This document is the ADAM DNA Deployment Specification for {self.company_name}, '
            f'targeting {details.get("display", platform)}. It was auto-generated from the '
            f'completed ADAM DNA Questionnaire and contains all infrastructure, agent configuration, '
            f'governance parameters, and deployment instructions needed to provision the full ADAM '
            f'autonomous operating model.'
        )
        doc.add_paragraph(
            f'ADAM Version: 1.1 | DNA Questionnaire Version: 1.0 | '
            f'Total Agents: {sum(len(c["agents"]) for c in AGENT_CLASSES.values())} | '
            f'Platform: {details.get("display", platform)}'
        )

        # Section 2: Company Doctrine Summary
        doc.add_heading('2. Company Doctrine Summary', level=1)
        self._add_doctrine_section(doc)

        # Section 3: Platform Architecture
        doc.add_heading('3. Platform Architecture', level=1)
        self._add_platform_architecture(doc, platform, details)

        # Section 4: CORE Engine Configuration
        doc.add_heading('4. CORE Engine Configuration', level=1)
        self._add_core_engine_section(doc, details)

        # Section 5: BOSS Scoring Configuration
        doc.add_heading('5. BOSS Scoring Configuration', level=1)
        self._add_boss_section(doc)

        # Section 6: Agent Mesh Deployment
        doc.add_heading('6. Agent Mesh Deployment', level=1)
        self._add_agent_mesh_section(doc, details)

        # Section 7: Governance & Exception Economy
        doc.add_heading('7. Governance & Exception Economy', level=1)
        self._add_governance_section(doc)

        # Section 8: Flight Recorder & Evidence Architecture
        doc.add_heading('8. Flight Recorder & Evidence Architecture', level=1)
        self._add_evidence_section(doc, details)

        # Section 9: Security & Resilience
        doc.add_heading('9. Security & Resilience', level=1)
        self._add_security_section(doc, platform)

        # Section 10: Deployment Procedure
        doc.add_heading('10. Deployment Procedure', level=1)
        self._add_deployment_procedure(doc, platform)

        # Appendix A: Agent Registry
        doc.add_page_break()
        doc.add_heading('Appendix A: Complete Agent Registry', level=1)
        self._add_agent_registry_appendix(doc)

        # Appendix B: Generated Files
        doc.add_heading('Appendix B: Generated Artifacts', level=1)
        doc.add_paragraph(
            'This deployment specification is accompanied by machine-readable artifacts '
            'in the output directory, including Infrastructure-as-Code templates (Terraform, '
            'Bicep, CloudFormation, Helm), JSON/YAML configuration bundles, CORE Graph seed '
            'data, BOSS scoring policies (OPA/Rego), and agent registry files.'
        )

        # Save
        filename = f"ADAM_DNA_Deployment_Spec_{platform.replace('-', '_')}_{self.company_name.replace(' ', '_')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        doc.save(filepath)
        return filepath

    def _add_title_page(self, doc, platform, details):
        for _ in range(4):
            doc.add_paragraph()

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('ADAM')
        run.bold = True
        run.font.size = Pt(36)
        run.font.color.rgb = RGBColor(0, 51, 102)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Autonomy Doctrine & Architecture Model')
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 51, 102)

        doc.add_paragraph()

        spec_title = doc.add_paragraph()
        spec_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = spec_title.add_run('DNA Deployment Specification')
        run.bold = True
        run.font.size = Pt(24)

        platform_name = doc.add_paragraph()
        platform_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = platform_name.add_run(details.get("display", platform))
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(102, 102, 102)

        doc.add_paragraph()
        doc.add_paragraph()

        company = doc.add_paragraph()
        company.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = company.add_run(f'Company: {self.company_name}')
        run.bold = True
        run.font.size = Pt(14)

        version = doc.add_paragraph()
        version.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = version.add_run('Version 1.0 | March 2026')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(102, 102, 102)

        gen = doc.add_paragraph()
        gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = gen.add_run('Generated by ADAM DNA Deployment Tool')
        run.italic = True
        run.font.size = Pt(10)

    def _add_doctrine_section(self, doc):
        gov = self.dna.get("governance", {})

        doc.add_heading('2.1 Mission', level=2)
        doc.add_paragraph(gov.get("mission", "[From DNA Questionnaire Section 1.1.2]"))

        doc.add_heading('2.2 Vision', level=2)
        doc.add_paragraph(gov.get("vision", "[From DNA Questionnaire Section 1.1.3]"))

        doc.add_heading('2.3 Permanent Principles', level=2)
        principles = gov.get("principles", [])
        if principles:
            for p in principles:
                doc.add_paragraph(p, style='List Bullet')
        else:
            doc.add_paragraph("[From DNA Questionnaire Section 1.1.4]")

        doc.add_heading('2.4 Human Governance Structure', level=2)
        doc.add_paragraph(
            'ADAM operates with a 5-Director Constitution. Directors do not manage workflows — '
            'they govern intent, constraints, and exceptions. They interact with ADAM exclusively '
            'through the Intent Interpretation Agent, Trust Gateway Agent, and Explain-Back Agent.'
        )
        directors_text = gov.get("directors", "")
        if directors_text:
            doc.add_paragraph(str(directors_text)[:500] + "..." if len(str(directors_text)) > 500 else str(directors_text))

    def _add_platform_architecture(self, doc, platform, details):
        doc.add_heading('3.1 Platform Service Mapping', level=2)

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = 'ADAM Component'
        hdr[1].text = f'{details.get("display", platform)} Service'

        mappings = [
            ('CORE Graph Database', details.get('graph_db', 'N/A')),
            ('Container Orchestration', details.get('container', 'N/A')),
            ('AI/LLM Services', details.get('ai', 'N/A')),
            ('Flight Recorder (Evidence)', details.get('evidence', 'N/A')),
            ('Cryptographic Vault', details.get('vault', 'N/A')),
            ('Messaging/Orchestration', details.get('messaging', 'N/A')),
            ('Monitoring/Observability', details.get('monitoring', 'N/A')),
            ('BOSS Score History', details.get('scoring', 'N/A')),
        ]

        for component, service in mappings:
            row = table.add_row().cells
            row[0].text = component
            row[1].text = service

        doc.add_heading('3.2 Architecture Principles', level=2)
        doc.add_paragraph(
            'Sovereignty-First: The governance layer (CORE Engine, policy enforcement, BOSS scoring, '
            'Flight Recorder, Cryptographic Vault) must remain under jurisdictional control. '
            'Control planes are sovereign; data planes can be distributed.'
        )

    def _add_core_engine_section(self, doc, details):
        doc.add_heading('4.1 Graph Database', level=2)
        doc.add_paragraph(f'Primary: {details.get("graph_db", "N/A")}')
        doc.add_paragraph(
            'The CORE Graph stores the company\'s Culture, Objectives, Rules, and Expectations '
            'as interconnected machine-checkable context. It contains 11 vertex types and 10 edge types.'
        )

        doc.add_heading('4.2 Vertex Types', level=2)
        vertices = ['Doctrine (root, immutable)', 'Culture', 'Objectives', 'Rules', 'Expectations',
                     'Financials Subgraph', 'Rights & Licensing Subgraph', 'Customer & Reputation Subgraph',
                     'Regulatory & Jurisdiction Subgraph', 'Strategy Drift Subgraph', 'Intent Objects']
        for v in vertices:
            doc.add_paragraph(v, style='List Bullet')

    def _add_boss_section(self, doc):
        boss = self.dna.get("boss_config", {})
        dims = boss.get("dimensions", {})

        doc.add_heading('5.1 BOSS Dimensions & Weights (v3.2)', level=2)
        doc.add_paragraph(
            'The BOSS (Business Operations Sovereignty Score) v3.2 evaluates every proposed action '
            'across the canonical seven dimensions. Composite Score C = (Sum(S_d x W_d)) / 24.0, '
            'where weights follow the Priority Tier model: Top=5.0 (Security Impact), '
            'Very High=4.0 (Sovereignty Action, Financial Exposure), High=3.0 (Regulatory Impact, '
            'Reputational Risk, Rights Certainty), Medium=2.0 (Doctrinal Alignment). '
            'Critical Dimension Override: if max(S_d) > 75, C = max(C, max(S_d) - 10). '
            'Non-Idempotent Penalty: +15 flat additive (not a multiplier).'
        )

        priority_tiers = {
            "security_impact": "Top",
            "sovereignty_action": "Very High",
            "financial_exposure": "Very High",
            "regulatory_impact": "High",
            "reputational_risk": "High",
            "rights_certainty": "High",
            "doctrinal_alignment": "Medium",
        }
        if dims:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text = 'Dimension'
            hdr[1].text = 'Weight'
            hdr[2].text = 'Priority Tier'
            for name, weight in dims.items():
                row = table.add_row().cells
                row[0].text = name.replace('_', ' ').title()
                row[1].text = str(weight)
                row[2].text = priority_tiers.get(name, "")

        doc.add_heading('5.2 Routing Thresholds', level=2)
        thresholds = [
            ('0-10', 'SOAP', 'Safe & Optimum Autonomous Performance - full autonomous execution, minimal logging'),
            ('11-30', 'MODERATE', 'Autonomous execution with enhanced logging, post-hoc review'),
            ('31-50', 'ELEVATED', 'Structured exception packet, relevant Domain Governor Agent reviews'),
            ('51-75', 'HIGH', 'Mandatory human director review, 4-hour max response'),
            ('76-100', 'OHSHAT', 'Operational Hell, Send Humans Act Today! - CEO + all directors, automatic safe-mode'),
        ]
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = 'Score Range'
        hdr[1].text = 'Tier'
        hdr[2].text = 'Action'
        for score, tier, action in thresholds:
            row = table.add_row().cells
            row[0].text = score
            row[1].text = tier
            row[2].text = action

    def _add_agent_mesh_section(self, doc, details):
        doc.add_heading('6.1 Agent Mesh Summary', level=2)

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = 'Agent Class'
        hdr[1].text = 'Count'
        hdr[2].text = 'Description'

        for class_key, class_data in AGENT_CLASSES.items():
            row = table.add_row().cells
            row[0].text = class_key.replace('_', ' ').title()
            row[1].text = str(len(class_data['agents']))
            row[2].text = class_data['description']

        total_row = table.add_row().cells
        total_row[0].text = 'TOTAL'
        run = total_row[0].paragraphs[0].runs[0]
        run.bold = True
        total_row[1].text = str(sum(len(c['agents']) for c in AGENT_CLASSES.values()))
        total_row[2].text = ''

        doc.add_heading('6.2 Container Orchestration', level=2)
        doc.add_paragraph(f'Platform: {details.get("container", "N/A")}')

    def _add_governance_section(self, doc):
        doc.add_heading('7.1 Exception Economy', level=2)
        doc.add_paragraph(
            'ADAM operates on the Exception Economy principle: autonomy is the default state, '
            'and human involvement is earned by consequence. Directors manage exceptions, not workflows.'
        )

        doc.add_heading('7.2 Intent Object Pipeline', level=2)
        doc.add_paragraph(
            'The six-stage intent-to-execution pipeline: '
            '1) Intent Interpretation, 2) Trust Gateway validation, '
            '3) Governor Agent governance evaluation, 4) BOSS scoring, '
            '5) Orchestration decomposition, 6) Evidence-first execution.'
        )

        doc.add_heading('7.3 Doctrine Conflict Arbitration', level=2)
        doc.add_paragraph(
            'When irreconcilable tensions emerge between doctrine elements, ADAM formalizes them '
            'as governance events routed to human directors. ADAM never silently resolves doctrine conflicts.'
        )

    def _add_evidence_section(self, doc, details):
        doc.add_paragraph(f'Evidence Store: {details.get("evidence", "N/A")}')
        doc.add_paragraph(
            'Every action produces governance-grade evidence by construction, not by after-the-fact '
            'reconstruction. All entries are cryptographically signed, hash-chained, and immutable. '
            'Audit = playback, not archaeology.'
        )

    def _add_security_section(self, doc, platform):
        infra = self.dna.get("infrastructure", {})
        doc.add_heading('9.1 Resilience Tiers', level=2)
        if infra.get("resilience_tiers"):
            doc.add_paragraph(str(infra["resilience_tiers"])[:800])

        doc.add_heading('9.2 Security Posture', level=2)
        if infra.get("security_posture"):
            doc.add_paragraph(str(infra["security_posture"])[:800])

    def _add_deployment_procedure(self, doc, platform):
        doc.add_paragraph(
            'The ADAM deployment follows a 7-phase bootstrap procedure with validation gates, '
            'cost estimation, and rollback capabilities at each phase.'
        )
        phases = [
            'Phase 1: Infrastructure provisioning (networking, identity, storage)',
            'Phase 2: CORE Engine deployment (graph database, policy engine)',
            'Phase 3: Governance plane activation (BOSS scoring, Flight Recorder, Crypto Vault)',
            'Phase 4: Agent mesh deployment (81+ agents across 7 classes)',
            'Phase 5: CORE Graph seeding (doctrine, culture, objectives, rules from DNA Questionnaire)',
            'Phase 6: Digital Twin initialization and calibration',
            'Phase 7: Validation, human interface testing, and go-live',
        ]
        for phase in phases:
            doc.add_paragraph(phase, style='List Number')

    def _add_agent_registry_appendix(self, doc):
        for class_key, class_data in AGENT_CLASSES.items():
            doc.add_heading(f'{class_key.replace("_", " ").title()} ({len(class_data["agents"])} agents)', level=2)
            doc.add_paragraph(class_data["description"])

            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text = 'Agent ID'
            hdr[1].text = 'Name'
            hdr[2].text = 'vCPUs'
            hdr[3].text = 'RAM (GB)'

            for agent in class_data["agents"]:
                row = table.add_row().cells
                row[0].text = agent["id"]
                row[1].text = agent["name"]
                row[2].text = str(agent["resources"]["vcpus"])
                row[3].text = str(agent["resources"]["ram_gb"])

            doc.add_paragraph()
